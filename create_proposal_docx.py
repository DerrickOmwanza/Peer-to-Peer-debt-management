#!/usr/bin/env python3
"""
Create comprehensive Word document for Peer-to-Peer Debt Management System Proposal
Based on detailed proposal structure provided
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_styled(doc, text, level):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.style = 'Heading 1'
        for run in heading.runs:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def create_table_from_data(doc, headers, data):
    """Create a table from headers and data"""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
    
    # Data rows
    for row_idx, row_data in enumerate(data, 1):
        row_cells = table.rows[row_idx].cells
        for col_idx, value in enumerate(row_data):
            row_cells[col_idx].text = str(value)
    
    return table

def create_word_document():
    """Create comprehensive Word document"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ============ COVER PAGE ============
    title = doc.add_heading('PEER-TO-PEER DEBT MANAGEMENT SYSTEM', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Project Proposal')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Cover page details
    cover_details = [
        ('Project Name:', 'Peer-to-Peer Debt Management System'),
        ('Author:', 'Derrick Omwanza Atandi'),
        ('Version:', '1.0'),
        ('Date:', '03 February 2026'),
        ('Status:', 'For Implementation'),
    ]
    
    for label, value in cover_details:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        p.add_run(f' {value}')
    
    doc.add_page_break()
    
    # ============ ONE-LINE SUMMARY ============
    add_heading_styled(doc, 'One-Line Summary', 1)
    doc.add_paragraph(
        'A standalone digital platform that formalizes peer-to-peer loans and enforces '
        'repayments by linking to M-PESA through user initiated prompt transfers.'
    )
    
    doc.add_page_break()
    
    # ============ EXECUTIVE SUMMARY ============
    add_heading_styled(doc, 'Executive Summary', 1)
    doc.add_paragraph(
        'This proposal describes a standalone Peer-to-Peer Debt Management System that enables '
        'individuals and groups to create legally auditable loan agreements, automate repayments '
        'via M-PESA user initiated prompts, and assess borrower reliability using explainable '
        'ML scoring. The platform reduces defaults, builds trust in community lending, and offers '
        'monetization through transaction fees, premium analytics, and licensing to SACCOs and Chamas. '
        'The design prioritizes clear consent, auditability, regulatory compliance, and conservative '
        'penalty policies.'
    )
    
    doc.add_page_break()
    
    # ============ PROBLEM STATEMENT ============
    add_heading_styled(doc, 'Problem Statement', 1)
    doc.add_paragraph(
        'Informal lending in Kenya is widespread but suffers from high default rates and disputes '
        'because agreements are verbal or informal. Lenders lack reliable enforcement tools and '
        'borrowers lack structured ways to commit to repayment terms. Existing M-PESA services focus '
        'on institutional credit and do not provide peer to peer debt management features. This gap '
        'discourages community lending and creates financial friction.'
    )
    
    doc.add_page_break()
    
    # ============ PROPOSED SOLUTION ============
    add_heading_styled(doc, 'Proposed Solution', 1)
    doc.add_paragraph(
        'A standalone digital platform where users register, create loan agreements with clause by '
        'clause consent, and authorize repayment rules. The system does not custody funds; instead '
        'it uses M-PESA user initiated prompt transfers to move money when repayments are due or when '
        'incoming funds meet configured thresholds. The platform stores immutable ledger entries, '
        'generates signed agreement PDFs for evidence, and provides ML based advisory scores to lenders. '
        'Dispute and reversal workflows protect users from erroneous deductions.'
    )
    
    doc.add_page_break()
    
    # ============ KEY FEATURES ============
    add_heading_styled(doc, 'Key Features', 1)
    
    features = [
        'Digital Agreements – clause by clause consent and electronic signature',
        'M-PESA Prompt Repayments – user initiated transfer prompts for repayments and disbursements',
        'Debt Ledger – immutable, timestamped records and exportable evidence',
        'Repayment Triggers – threshold and date based triggers for repayments',
        'Configurable Repayment Methods – fixed amounts, percentage of incoming funds, scheduled installments',
        'Penalty Rules – configurable late fees with legal caps and grace periods',
        'ML Risk Scoring – explainable borrower reliability scores shown as advisory to lenders',
        'Notifications – SMS and in app alerts for all key events',
        'Dispute and Reversal Flows – user dispute submission, evidence capture, and manual review',
        'KYC and Fraud Controls – identity verification and transaction monitoring',
        'Dashboard – single user view for loans issued and loans owed',
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # ============ SYSTEM ARCHITECTURE ============
    add_heading_styled(doc, 'System Architecture Overview', 1)
    
    add_heading_styled(doc, 'Components and Interactions', 2)
    
    components = [
        ('User Interface', 'Web and mobile app for smartphones; USSD flow for feature phones'),
        ('Consent and Agreement Engine', 'Presents terms, captures clause ticks, records electronic signature or USSD acceptance'),
        ('Debt Ledger', 'Secure database with immutable entries, transaction references, and audit trail'),
        ('Trigger Engine', 'Monitors configured triggers and issues M-PESA prompt requests to users'),
        ('M-PESA Integration Layer', 'Orchestrates user initiated prompt flows and receives callbacks/receipts'),
        ('Notification Service', 'Sends SMS and in app notifications for approvals, repayments, disputes'),
        ('ML Risk Scoring Module', 'Generates explainable scores from platform behavior and optional alternative data'),
        ('Admin and Operations Console', 'Manual review, dispute resolution, reporting, and compliance tools'),
    ]
    
    for component, description in components:
        p = doc.add_paragraph()
        run = p.add_run(component)
        run.bold = True
        p.add_run(f' – {description}')
    
    add_heading_styled(doc, 'Interaction Summary', 2)
    doc.add_paragraph(
        'User creates agreement → Agreement stored in ledger → Lender approves → Funds disbursed via M-PESA prompt → '
        'Repayment triggers monitored → When triggered, system issues M-PESA prompt to debtor to transfer repayment → '
        'Callback updates ledger and notifies parties.'
    )
    
    doc.add_page_break()
    
    # ============ USER FLOW ============
    add_heading_styled(doc, 'User Flow – Loan Request to Repayment', 1)
    
    add_heading_styled(doc, '1. Loan Initiation', 2)
    loan_init = [
        'User logs in or registers',
        'User selects Create Loan Request',
        'User fills: loan amount, lender phone number, repayment start date, repayment method, threshold, interest option, penalty rules',
        'System displays clause by clause terms',
        'User ticks each clause and signs electronically or accepts via USSD tick',
        'User enters M-PESA PIN to authorize prompt usage and submits request',
        'System stores agreement and sends lender a prompt',
    ]
    for step in loan_init:
        doc.add_paragraph(step, style='List Number')
    
    add_heading_styled(doc, '2. Lender Approval', 2)
    lender_approval = [
        'Lender receives SMS/app prompt with loan details and borrower ML score',
        'Lender accepts or declines',
        'If accepted, lender enters M-PESA PIN to send funds via M-PESA prompt',
        'On successful transfer, ledger records disbursement and initializes balance',
    ]
    for step in lender_approval:
        doc.add_paragraph(step, style='List Number')
    
    add_heading_styled(doc, '3. Repayment Triggering – Two Pathways', 2)
    
    add_heading_styled(doc, 'Pathway A: MVP Flow (Daraja STK Push – Current Implementation)', 3)
    doc.add_paragraph(
        'Achievable today using Safaricom Daraja APIs. Ensures compliance with current M-PESA rules and user consent.'
    )
    repayment_trigger_mvp = [
        'On repayment start date, Trigger Engine monitors incoming transactions to debtor account',
        'When incoming transaction ≥ configured threshold (e.g., Ksh 100), system detects receipt',
        'System issues M-PESA STK Push prompt to debtor with: Repayment amount, Lender phone, Loan reference, and Clear consent text',
        'Debtor receives on-device STK popup: "Ksh 200 repayment due for loan LN-12345 to 07YYYYYYYY. Confirm and enter M-PESA PIN to transfer."',
        'Debtor confirms by entering M-PESA PIN (demonstrates active consent and authorization)',
        'M-PESA processes the transfer and sends callback to platform with transaction reference',
        'On successful callback, Ledger updates balance, records transaction, and notifies both parties with receipt',
    ]
    for step in repayment_trigger_mvp:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph()
    doc.add_paragraph('Key Advantages of MVP Flow:').runs[0].bold = True
    mvp_advantages = [
        'Compliant with current Safaricom policies (no auto-deduction)',
        'User has explicit control and can review repayment before PIN entry',
        'Reduces borrower distrust and friction',
        'Achievable within 12-week MVP timeline',
        'Demonstrates platform viability for regulatory review',
    ]
    for adv in mvp_advantages:
        doc.add_paragraph(adv, style='List Bullet')
    
    doc.add_paragraph()
    add_heading_styled(doc, 'Pathway B: Future Flow (Fuliza-Style Auto-Deduction – Partnership Scenario)', 3)
    doc.add_paragraph(
        'Requires formal partnership with Safaricom and regulatory approval (CBK clearance). Enables automated enforcement '
        'and prevents repayment evasion.'
    )
    repayment_trigger_future = [
        'On repayment start date, Trigger Engine monitors incoming transactions to debtor account',
        'When incoming transaction ≥ configured threshold, system detects receipt in real-time',
        'System calculates repayment amount (fixed amount or percentage of incoming funds)',
        'System immediately deducts repayment amount directly from debtor wallet (no STK prompt required)',
        'Debtor receives SMS notification: "Ksh 200 auto-deducted for loan LN-12345. Balance: Ksh 4800. Dispute within 24h: reply DISPUTE"',
        'Lender receives SMS notification: "Ksh 200 received from 07XXXXXXXX. Balance: Ksh 4800."',
        'If debtor disputes within 24h, system initiates reversal and manual review workflow',
        'Transaction locked in ledger with auto-deduction timestamp and evidence for auditability',
    ]
    for step in repayment_trigger_future:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph()
    doc.add_paragraph('Key Advantages of Fuliza-Style Flow:').runs[0].bold = True
    future_advantages = [
        'Mimics Fuliza overdraft recovery (proven enforcement model)',
        'Prevents repayment evasion (debtor cannot "forget" to repay)',
        'Faster settlement and reduced administrative overhead',
        'Scales to high-volume lending without manual intervention',
        'Higher repayment completion rates (industry benchmark: 95%+ vs. 50-60% manual)',
        'Dispute-reversal mechanism protects against erroneous deductions',
    ]
    for adv in future_advantages:
        doc.add_paragraph(adv, style='List Bullet')
    
    add_heading_styled(doc, '4. Missed Payment and Penalty Application', 2)
    missed_payment = [
        'If repayment not completed by due date, system applies configured penalty rules after grace period',
        'Penalties are capped and displayed in agreement',
        'Notifications sent to debtor and lender',
    ]
    for step in missed_payment:
        doc.add_paragraph(step, style='List Number')
    
    add_heading_styled(doc, '5. Dispute and Reversal', 2)
    dispute = [
        'Either party can raise dispute via dashboard or SMS',
        'System captures evidence and routes to manual review',
        'If reversal is warranted, operations team initiates reversal and ledger is updated',
    ]
    for step in dispute:
        doc.add_paragraph(step, style='List Number')
    
    add_heading_styled(doc, '6. Loan Closure', 2)
    closure = [
        'When balance reaches zero, system marks loan closed and issues final receipts and signed agreement PDF',
    ]
    for step in closure:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # ============ FUNCTIONAL REQUIREMENTS ============
    add_heading_styled(doc, 'Detailed Functional Requirements', 1)
    
    add_heading_styled(doc, 'User Roles', 2)
    roles = [
        'Borrower – creates loan requests and authorizes repayments',
        'Lender – approves requests and disburses funds',
        'Admin – manages disputes, compliance, and operations',
        'System – automated components: Trigger Engine, Notification Service, ML Module',
    ]
    for role in roles:
        doc.add_paragraph(role, style='List Bullet')
    
    add_heading_styled(doc, 'User Stories', 2)
    stories = [
        'As a borrower I want to create a loan request and capture consent so that I have a legally auditable agreement',
        'As a lender I want to see borrower reliability and accept or decline a loan request',
        'As a user I want repayments to be automatically prompted when I receive funds above a threshold',
        'As an admin I want to review disputes and reverse erroneous transactions',
    ]
    for i, story in enumerate(stories, 1):
        doc.add_paragraph(story, style='List Number')
    
    add_heading_styled(doc, 'Workflows and Expected System Responses', 2)
    
    workflows = [
        ('Create Loan Request', 'System validates inputs, stores agreement, returns confirmation ID'),
        ('Approve Loan', 'System sends M-PESA prompt to lender; on success, records disbursement and notifies borrower'),
        ('Trigger Repayment', 'System detects incoming transaction, issues prompt to debtor; on callback, updates ledger and notifies both parties'),
        ('Record Repayment', 'System logs transaction reference, updates balance, recalculates interest and penalties'),
        ('Dispute Handling', 'System logs dispute, locks affected ledger entries, notifies operations, and provides evidence package'),
    ]
    
    for workflow, response in workflows:
        p = doc.add_paragraph()
        run = p.add_run(workflow)
        run.bold = True
        p.add_run(f' – {response}')
    
    doc.add_page_break()
    
    # ============ NON-FUNCTIONAL REQUIREMENTS ============
    add_heading_styled(doc, 'Non-Functional Requirements', 1)
    
    nfr_items = [
        ('Performance', 'System must handle 1,000 concurrent users in pilot and scale to 100,000 monthly active users'),
        ('Scalability', 'Microservices architecture with autoscaling for Trigger Engine and Notification Service'),
        ('Availability', '99.9% uptime for core services; 99.95% for ledger read operations'),
        ('Latency', 'API responses under 300 ms for core endpoints under normal load'),
        ('Reliability', 'Durable ledger with transactional integrity and idempotent operations'),
        ('Auditability', 'Immutable logs for all consent, approvals, and transactions with exportable evidence'),
        ('Data Retention', 'Retain agreements and transaction records for minimum 7 years or as required by law'),
        ('Encryption', 'TLS 1.2+ in transit; AES 256 at rest'),
        ('Backups', 'Daily backups with point in time recovery for 30 days'),
    ]
    
    for nfr, description in nfr_items:
        p = doc.add_paragraph()
        run = p.add_run(nfr)
        run.bold = True
        p.add_run(f' – {description}')
    
    doc.add_page_break()
    
    # ============ SECURITY AND PRIVACY ============
    add_heading_styled(doc, 'Security and Privacy', 1)
    
    security_items = [
        ('Authentication', 'Multi factor authentication for web and mobile; USSD flows use phone verification'),
        ('Authorization', 'Role based access control with least privilege'),
        ('PIN Handling', 'M-PESA PIN is never stored; only used to initiate M-PESA prompt flows on the user device'),
        ('Storage Encryption', 'Sensitive fields encrypted with customer specific keys'),
        ('Key Management', 'Use HSM or cloud KMS for key storage and rotation'),
        ('Secure Logging', 'Logs redact sensitive data and include tamper evident hashes'),
        ('Fraud Protection', 'Transaction anomaly detection, velocity checks, and manual review triggers'),
        ('KYC Requirements', 'Identity verification for users above configured thresholds; support for national ID capture and verification'),
        ('Consent Capture', 'Clause level consent stored with timestamps and device metadata'),
        ('Data Minimization', 'Collect only necessary personal data and provide deletion workflows per Data Protection Act'),
        ('Retention and Deletion', 'Retention policy aligned with legal requirements; deletion requests processed with audit trail'),
    ]
    
    for security, description in security_items:
        p = doc.add_paragraph()
        run = p.add_run(security)
        run.bold = True
        p.add_run(f' – {description}')
    
    doc.add_page_break()
    
    # ============ REGULATORY AND LEGAL ============
    add_heading_styled(doc, 'Regulatory and Legal Considerations', 1)
    
    regulatory_items = [
        ('CBK Digital Credit Rules', 'Confirm whether platform qualifies as a Digital Credit Provider and obtain necessary licensing or partner with a licensed entity'),
        ('Data Protection Act Compliance', 'Implement lawful basis for processing, data subject rights, DPIA, and data breach procedures'),
        ('Consumer Protection', 'Transparent disclosure of fees, interest, penalties, and dispute resolution mechanisms'),
        ('Interest and Penalty Legality', 'Avoid punitive daily interest rates such as 10% per day without legal validation. Use conservative, legally vetted penalty structures with caps and grace periods'),
        ('Evidence Admissibility', 'Signed agreement PDFs must include timestamps, device metadata, and transaction references to support court admissibility'),
    ]
    
    for reg, description in regulatory_items:
        p = doc.add_paragraph()
        run = p.add_run(reg)
        run.bold = True
        p.add_run(f' – {description}')
    
    add_heading_styled(doc, 'Suggested Legal Review Items', 2)
    legal_items = [
        'Licensing requirements',
        'Permissible penalty rates',
        'Enforceability of electronic signatures',
        'Partnership terms with Safaricom',
    ]
    for item in legal_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_styled(doc, 'Safaricom Partnership Considerations', 2)
    doc.add_paragraph(
        'Negotiate API access or formalize user initiated prompt flows; define SLAs, callback formats, '
        'and dispute handling responsibilities.'
    )
    
    doc.add_page_break()
    
    # ============ ML RISK SCORING ============
    add_heading_styled(doc, 'ML Risk Scoring Module', 1)
    
    add_heading_styled(doc, 'Purpose', 2)
    doc.add_paragraph(
        'Provide lenders with an explainable advisory score indicating borrower likelihood to repay on time.'
    )
    
    add_heading_styled(doc, 'Inputs and Features', 2)
    scoring_inputs = [
        'On platform repayment history',
        'Frequency and size of loans',
        'Timeliness of past repayments',
        'Dispute history',
        'Optional alternative data with consent: mobile money transaction patterns, airtime purchases, utility payments',
    ]
    for input_item in scoring_inputs:
        doc.add_paragraph(input_item, style='List Bullet')
    
    add_heading_styled(doc, 'Model Types', 2)
    doc.add_paragraph(
        'Start with interpretable models: logistic regression or gradient boosted trees with SHAP explainability. '
        'Consider neural models only after sufficient data and explainability controls.'
    )
    
    add_heading_styled(doc, 'Explainability Requirements', 2)
    explainability = [
        'Provide top 3 features influencing each score',
        'Display score band and confidence interval',
        'Scores are advisory only and must not be the sole basis for automated denial',
    ]
    for item in explainability:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_styled(doc, 'Fairness and Bias Mitigation', 2)
    fairness = [
        'Monitor for demographic bias',
        'Use fairness metrics and regular audits',
        'Allow users to contest scores and provide remediation paths',
    ]
    for item in fairness:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_styled(doc, 'Training and Retraining', 2)
    doc.add_paragraph(
        'Retrain monthly during pilot; quarterly in production. Maintain training data lineage and versioning.'
    )
    
    add_heading_styled(doc, 'Performance Metrics', 2)
    doc.add_paragraph(
        'AUC, precision at top deciles, calibration, and false positive/negative rates.'
    )
    
    add_heading_styled(doc, 'Presentation to Lenders', 2)
    doc.add_paragraph(
        'Show score as a band (Low, Medium, High risk) with numeric value and top contributing factors. '
        'Include advisory text: "Score is advisory only. Consider additional context."'
    )
    
    doc.add_page_break()
    
    # ============ REPAYMENT AND PENALTY RULES ============
    add_heading_styled(doc, 'Repayment and Penalty Rules', 1)
    
    add_heading_styled(doc, 'Configurable Repayment Methods', 2)
    methods = [
        'Fixed Installment – set amount per period',
        'Percentage of Incoming Funds – e.g., 20% of incoming transactions above threshold',
        'One time Due Date – full balance due on specified date',
    ]
    for method in methods:
        doc.add_paragraph(method, style='List Bullet')
    
    add_heading_styled(doc, 'Threshold Rules', 2)
    doc.add_paragraph(
        'Minimum incoming transaction amount to trigger repayment, e.g., Ksh 100.'
    )
    
    add_heading_styled(doc, 'Sample Calculations', 2)
    doc.add_paragraph(
        'Fixed: Loan Ksh 5,000; repayment Ksh 200 per incoming trigger; 25 successful triggers to repay principal.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Percentage: Loan Ksh 5,000; 10% of incoming funds; if incoming Ksh 2,000 then repayment Ksh 200.',
        style='List Bullet'
    )
    
    add_heading_styled(doc, 'Penalty Application Rules', 2)
    penalty_rules = [
        'Grace Period – configurable, default 7 days',
        'Penalty Rate – recommend conservative monthly rate or flat late fee; avoid extreme daily compounding',
        'Caps – total penalties capped at a percentage of principal, e.g., 50% of principal',
        'Transparency – all penalties displayed in agreement and recalculated daily with audit trail',
    ]
    for rule in penalty_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_page_break()
    
    # ============ DISPUTE AND REVERSAL ============
    add_heading_styled(doc, 'Dispute, Reversal, and Exception Handling', 1)
    
    add_heading_styled(doc, 'User-Initiated Disputes', 2)
    doc.add_paragraph('Users submit dispute with reason and supporting evidence.')
    doc.add_paragraph('System assigns dispute ID and acknowledges receipt.')
    
    add_heading_styled(doc, 'Automated Reversals', 2)
    doc.add_paragraph(
        'If system detects duplicate or erroneous deduction, auto flag and initiate reversal workflow with operations approval.'
    )
    
    add_heading_styled(doc, 'Manual Review Workflow', 2)
    doc.add_paragraph('Operations console shows dispute, evidence, transaction references, and agreement PDF.')
    doc.add_paragraph('SLA: initial response within 24 hours; resolution within 7 business days for pilot.')
    
    add_heading_styled(doc, 'Evidence Required', 2)
    evidence = [
        'Agreement PDF',
        'Transaction callback reference',
        'Timestamps',
        'User statements',
    ]
    for item in evidence:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_styled(doc, 'Notification Templates', 2)
    doc.add_paragraph(
        'Dispute Received: "We received your dispute ID 12345. We will respond within 24 hours."',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Reversal Completed: "Ksh 200 reversed to your account. Reference: R 98765."',
        style='List Bullet'
    )
    
    doc.add_page_break()
    
    # ============ SCALABILITY AND PATHWAY PROGRESSION ============
    add_heading_styled(doc, 'System Scalability: From MVP to Enterprise Scale', 1)
    
    add_heading_styled(doc, 'Two-Pathway Architecture Strategy', 2)
    doc.add_paragraph(
        'The system is architected to support two distinct repayment enforcement models, allowing progression from '
        'MVP to enterprise scale without redesign. Both pathways coexist in the same ledger and API layer.'
    )
    
    add_heading_styled(doc, 'Phase 1: MVP Scale (Daraja STK Push)', 2)
    mvp_scale = [
        'Timeline: Months 1-3 (current)',
        'Users: 500-1,000 borrowers and lenders',
        'Monthly loans: 500-2,000',
        'Repayment model: STK Push with PIN confirmation',
        'Infrastructure: Single-region deployment, standard database',
        'Regulatory status: No special licensing required (user-initiated transfers)',
        'Adoption risk: Low (user retains control)',
        'Completion rate: 60-70% (depends on user follow-through)',
    ]
    for item in mvp_scale:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('MVP Success Metrics:').runs[0].bold = True
    mvp_metrics = [
        'System demonstrates feasibility with Safaricom Daraja APIs',
        'Establishes repayment completion data and borrower behavior patterns',
        'Builds user trust through explicit consent and control',
        'Provides evidence for regulatory review and CBK submission',
        'Identifies technical and operational bottlenecks before scale',
    ]
    for metric in mvp_metrics:
        doc.add_paragraph(metric, style='List Bullet')
    
    add_heading_styled(doc, 'Phase 2: Enterprise Scale (Fuliza-Style Auto-Deduction)', 2)
    enterprise_scale = [
        'Timeline: Months 4-12 (post-MVP, pending partnership)',
        'Users: 50,000-100,000 borrowers and lenders',
        'Monthly loans: 50,000-100,000',
        'Repayment model: Auto-deduction from wallet (no PIN required)',
        'Infrastructure: Multi-region, real-time processing, distributed ledger',
        'Regulatory status: Formal Safaricom partnership and CBK approval required',
        'Adoption risk: Medium (requires regulatory clarity)',
        'Completion rate: 90-95% (automated, prevents evasion)',
    ]
    for item in enterprise_scale:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Enterprise Scaling Triggers:').runs[0].bold = True
    scaling_triggers = [
        'MVP achieves 70%+ repayment completion rate',
        'Safaricom partnership negotiation completed and API access granted',
        'CBK provides regulatory clearance for auto-deduction mechanism',
        'System demonstrates operational excellence (99.9% uptime, <300ms latency)',
        'User base reaches 10,000+ active lenders (minimum viable ecosystem)',
    ]
    for trigger in scaling_triggers:
        doc.add_paragraph(trigger, style='List Bullet')
    
    add_heading_styled(doc, 'Technical Compatibility: No Rework Required', 2)
    doc.add_paragraph(
        'Both pathways share identical data model, API contracts, and ledger schema. '
        'Migration from MVP to Fuliza-style flow requires:'
    )
    
    compatibility = [
        'Configuration change: repayment_method flag (STK_PUSH → AUTO_DEDUCT)',
        'Safaricom API integration: Add auto-deduction endpoint (Daraja partner-initiated transfer)',
        'Compliance layer: Add dispute-reversal callbacks for auto-deducted transactions',
        'Operational workflow: Enable 24-hour reversal window for disputed auto-deductions',
        'No database migration, no API redesign, no ledger restructuring',
    ]
    for item in compatibility:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ============ M-PESA INTEGRATION ============
    add_heading_styled(doc, 'Integration with M-PESA', 1)
    
    add_heading_styled(doc, 'Design Approach', 2)
    design_points = [
        'Platform remains standalone and uses M-PESA only via user initiated prompts',
        'No custody of funds',
        'Two primary flows: lender disbursement prompt and debtor repayment prompt',
    ]
    for point in design_points:
        doc.add_paragraph(point, style='List Bullet')
    
    add_heading_styled(doc, 'Sequence of Prompts and Confirmations', 2)
    sequence = [
        'Platform requests M-PESA prompt to lender with transaction details',
        'Lender receives prompt on device and confirms with M-PESA PIN',
        'M-PESA sends callback to platform with transaction reference',
        'Platform updates ledger and notifies borrower',
    ]
    for step in sequence:
        doc.add_paragraph(step, style='List Number')
    
    add_heading_styled(doc, 'Sample SMS/USSD Prompt Text and API Implementations', 2)
    
    add_heading_styled(doc, 'MVP Implementation (Daraja STK Push)', 3)
    doc.add_paragraph('Lender Disbursement Prompt:')
    doc.add_paragraph(
        '"You have a loan request to fund Ksh 5,000 to 07XXXXXXXX. Reply to confirm and enter your M-PESA PIN to complete."',
        style='List Bullet'
    )
    doc.add_paragraph('Debtor Repayment STK Push Prompt:')
    doc.add_paragraph(
        '"Ksh 200 repayment due for loan LN-12345 to 07YYYYYYYY. Confirm and enter your M-PESA PIN to transfer."',
        style='List Bullet'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('Daraja API Contract (STK Push):').runs[0].bold = True
    doc.add_paragraph('Endpoint: POST /mpesa/stkpush/v1/processrequest')
    doc.add_paragraph(
        'Request: { "BusinessShortCode": "174379", "Password": "[hashed]", "Timestamp": "20260203153000", '
        '"TransactionType": "CustomerPayBillOnline", "Amount": 200, "PartyA": "254700123456", "PartyB": "254708374149", '
        '"PhoneNumber": "254700123456", "CallBackURL": "https://api.example.com/v1/mpesa/stk-callback", '
        '"AccountReference": "LN-12345-Repayment", "TransactionDesc": "Loan repayment for LN-12345" }',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Response: { "ResponseCode": "0", "ResponseDescription": "success.", '
        '"MerchantRequestID": "29115-34620561-1", "CheckoutRequestID": "ws_CO_DMZ_123456789_07022025155500" }',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Callback on User Action: { "Body": { "stkCallback": { "MerchantRequestID": "...", "CheckoutRequestID": "...", '
        '"ResultCode": 0, "ResultDesc": "The service request has been processed successfully.", '
        '"CallbackMetadata": { "Item": [ { "Name": "Amount", "Value": 200 }, { "Name": "MpesaReceiptNumber", "Value": "MH12345A" }, '
        '{ "Name": "TransactionDate", "Value": "20260203153000" }, { "Name": "PhoneNumber", "Value": "254700123456" } ] } } } }',
        style='List Bullet'
    )
    
    doc.add_paragraph()
    add_heading_styled(doc, 'Future Implementation (Fuliza-Style Auto-Deduction)', 3)
    doc.add_paragraph('Debtor Receives Auto-Deduction Notification (No Prompt):')
    doc.add_paragraph(
        '"Ksh 200 auto-deducted for loan repayment LN-12345. New balance: Ksh 4800. Dispute within 24h: reply DISPUTE 98765."',
        style='List Bullet'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('Daraja Partner-Initiated Transfer API (Future):').runs[0].bold = True
    doc.add_paragraph('Endpoint: POST /mpesa/b2c/v3/payloads (with partner-initiated transfer flag)')
    doc.add_paragraph(
        'Request: { "InitiatorName": "[System service account]", "SecurityCredential": "[encrypted]", '
        '"CommandID": "PayLoanRepayment", "Amount": 200, "PartyA": "254708374149", "PartyB": "254700123456", '
        '"Remarks": "Loan LN-12345 repayment auto-deducted", "QueueTimeOutURL": "https://api.example.com/v1/mpesa/b2c-timeout", '
        '"ResultURL": "https://api.example.com/v1/mpesa/b2c-result", "Occasion": "LoanRepayment" }',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Response (Immediate): { "ConversationID": "AG_20260203_00001", "OriginatorConversationID": "123456789", '
        '"ResponseCode": "0", "ResponseDescription": "Accept the service request successfully." }',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Async Callback: { "Result": { "ResultCode": 0, "ResultDesc": "The service request has been processed successfully.", '
        '"ConversationID": "AG_20260203_00001", "OriginatorConversationID": "123456789", '
        '"TransactionAmount": 200, "B2CUtilityAccountAvailable": 0, "TransactionID": "MI123456789123", "TransactionReceipt": "MH12345B", '
        '"B2CChargesPaidAccountAvailable": 0, "B2BUtilityAccountAvailable": 0, "B2BChargesPaidAccountAvailable": 0 } }',
        style='List Bullet'
    )
    
    add_heading_styled(doc, 'Fallback When Prompt Fails', 2)
    fallback = [
        'Retry prompt up to 3 times',
        'Notify user to manually send funds and provide instructions',
        'If manual transfer occurs, user uploads transaction receipt and platform reconciles',
    ]
    for item in fallback:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ============ DATA MODEL AND API ============
    add_heading_styled(doc, 'Data Model and API Specifications', 1)
    
    add_heading_styled(doc, 'Key Tables and Fields', 2)
    
    tables_data = [
        ('Users', 'user_id (PK), phone, email, created_at'),
        ('Profiles', 'profile_id (PK), user_id (FK), name, national_id, kyc_status'),
        ('Loans', 'loan_id (PK), borrower_id (FK), lender_id (FK), principal, balance, status, created_at'),
        ('Agreements', 'agreement_id (PK), loan_id (FK), clauses (JSON), consent_timestamps (JSON), signature_meta (JSON)'),
        ('Transactions', 'transaction_id (PK), loan_id (FK), type, amount, mpesa_ref, status, timestamp'),
        ('Disputes', 'dispute_id (PK), transaction_id (FK), user_id (FK), status, evidence (JSON)'),
        ('Scores', 'score_id (PK), user_id (FK), score_value, features (JSON), created_at'),
    ]
    
    for table_name, fields in tables_data:
        p = doc.add_paragraph()
        run = p.add_run(f'{table_name}')
        run.bold = True
        p.add_run(f': {fields}')
    
    add_heading_styled(doc, 'Sample REST API Endpoints', 2)
    
    endpoints = [
        {
            'name': 'Create Loan Request',
            'method': 'POST /api/v1/loans',
            'request': '{ "borrower_id": "U1", "lender_phone": "07XXXX", "principal": 5000, "repayment_method": "fixed", "installment": 200, "start_date": "2026-03-01" }',
            'response': '{ "loan_id": "LN-12345", "status": "pending_lender_approval" }'
        },
        {
            'name': 'Approve Loan',
            'method': 'POST /api/v1/loans/LN-12345/approve',
            'request': '{ "lender_id": "U2" }',
            'response': '{ "status": "prompt_sent", "prompt_id": "P-98765" }'
        },
        {
            'name': 'Record Repayment',
            'method': 'POST /api/v1/transactions',
            'request': '{ "loan_id": "LN-12345", "amount": 200, "mpesa_ref": "M-PESA-TRX-001" }',
            'response': '{ "transaction_id": "T-555", "balance": 4800 }'
        },
        {
            'name': 'Query Ledger',
            'method': 'GET /api/v1/loans/LN-12345/ledger',
            'request': 'N/A',
            'response': '{ "loan_id": "LN-12345", "transactions": [ ... ] }'
        },
        {
            'name': 'Get Score',
            'method': 'GET /api/v1/users/U1/score',
            'request': 'N/A',
            'response': '{ "score": 0.72, "band": "Medium", "top_factors": ["late_payments", "loan_frequency"] }'
        },
    ]
    
    for endpoint in endpoints:
        p = doc.add_paragraph()
        run = p.add_run(endpoint['name'])
        run.bold = True
        p.add_run(f" – {endpoint['method']}")
        doc.add_paragraph(f"Request: {endpoint['request']}", style='List Bullet')
        doc.add_paragraph(f"Response: {endpoint['response']}", style='List Bullet')
    
    doc.add_page_break()
    
    # ============ UI/UX ============
    add_heading_styled(doc, 'UI and UX', 1)
    
    add_heading_styled(doc, 'Smartphone Screens', 2)
    screens = [
        'Dashboard – summary of loans owed and loans issued, quick action Create Loan',
        'Create Loan – form with fields, clause list with tick boxes, signature capture, submit button',
        'Loan Detail – agreement PDF, repayment schedule, transaction history, dispute button',
        'Notifications – in app and SMS with clear action items',
    ]
    for screen in screens:
        doc.add_paragraph(screen, style='List Bullet')
    
    add_heading_styled(doc, 'USSD and Feature Phone Flow', 2)
    doc.add_paragraph(
        'Start *XYZ# → Login via phone number → Select Create Loan → Enter lender phone → Enter amount → '
        'Accept terms by pressing 1 → Confirm.'
    )
    
    add_heading_styled(doc, 'Consent Screens', 2)
    doc.add_paragraph(
        'Clause list with short plain language statements and tick boxes. '
        'Final screen shows summary and signature or USSD acceptance.'
    )
    
    add_heading_styled(doc, 'Accessibility Considerations', 2)
    accessibility = [
        'High contrast UI, readable fonts, and simple language',
        'USSD flows for non smartphone users',
        'Local language support for key screens',
    ]
    for item in accessibility:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ============ TESTING AND ROLLOUT ============
    add_heading_styled(doc, 'Testing, QA, and Rollout Plan', 1)
    
    add_heading_styled(doc, 'Testing Strategy', 2)
    testing = [
        'Unit Tests – for business logic and ML pipelines',
        'Integration Tests – for M-PESA prompt flows and callbacks',
        'End to End Tests – simulating full loan lifecycle',
        'Security Tests – including penetration testing and code scanning',
    ]
    for test in testing:
        doc.add_paragraph(test, style='List Bullet')
    
    add_heading_styled(doc, 'Pilot Plan', 2)
    doc.add_paragraph('Pilot Size: 500–1,000 users in a defined region or partner SACCO')
    doc.add_paragraph('Duration: 3 months')
    doc.add_paragraph('Metrics to Monitor: repayment completion rate, dispute rate, prompt success rate, time to resolution, user NPS')
    
    add_heading_styled(doc, 'Rollback Criteria', 2)
    rollback = [
        'Prompt failure rate > 10% after retries',
        'Dispute rate > 5% of transactions',
        'Regulatory or legal objections from authorities',
    ]
    for criterion in rollback:
        doc.add_paragraph(criterion, style='List Bullet')
    
    add_heading_styled(doc, 'User Support Plan', 2)
    support = [
        'In app help center, SMS support, and phone support for disputes',
        'Operations team for manual reviews and reversals',
    ]
    for item in support:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ============ MVP SCOPE ============
    add_heading_styled(doc, 'MVP Scope and Roadmap', 1)
    
    add_heading_styled(doc, 'MVP Features', 2)
    mvp_features = [
        'User registration and basic KYC',
        'Create loan request with clause consent',
        'Lender approval prompt simulation',
        'Debt ledger and agreement PDF generation',
        'Trigger engine simulator for incoming transactions',
        'Notification system (SMS and email)',
        'Basic ML scoring prototype (rule based)',
    ]
    for feature in mvp_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_heading_styled(doc, 'Three-Phase Roadmap', 2)
    doc.add_paragraph(
        'Phase 1 (MVP): Core system with ledger, agreements, and basic triggers (Weeks 1-12)'
    )
    doc.add_paragraph(
        'Phase 2 (Enhancement): ML scoring, advanced analytics, multi-language support (Months 4-6)'
    )
    doc.add_paragraph(
        'Phase 3 (Scale): SACCO licensing, SIM Toolkit integration, institutional partnerships (Months 7-12)'
    )
    
    doc.add_page_break()
    
    # ============ CONCLUSION ============
    add_heading_styled(doc, 'Conclusion', 1)
    doc.add_paragraph(
        'This Peer-to-Peer Debt Management System addresses a clear market need by formalizing informal lending, '
        'improving repayment outcomes, and providing transparent, auditable agreements. The standalone design that '
        'links to M-PESA via user initiated prompts reduces custody risk while delivering a practical enforcement mechanism. '
        'To proceed, we recommend immediate legal review, a focused MVP build with a 3 month pilot, and outreach to potential '
        'partners such as SACCOs and Safaricom for pilot collaboration.'
    )
    
    doc.add_page_break()
    
    # ============ CALL TO ACTION ============
    add_heading_styled(doc, 'Call to Action', 1)
    doc.add_paragraph(
        'We invite partners, investors, and pilot collaborators to join a 3 month MVP pilot. '
        'Contact Derrick Omwanza Atandi to discuss partnership terms, pilot design, and legal next steps.'
    )
    
    doc.add_page_break()
    
    # ============ APPENDICES ============
    add_heading_styled(doc, 'Appendix A – Glossary', 1)
    
    glossary = [
        ('Agreement PDF', 'signed document capturing clauses and signatures'),
        ('DCP', 'Digital Credit Provider'),
        ('KYC', 'Know Your Customer'),
        ('ML', 'Machine Learning'),
        ('M-PESA Prompt', 'user initiated transfer request via M-PESA'),
        ('Trigger Engine', 'component that monitors incoming funds and issues prompts'),
    ]
    
    for term, definition in glossary:
        p = doc.add_paragraph()
        run = p.add_run(term)
        run.bold = True
        p.add_run(f' – {definition}')
    
    doc.add_page_break()
    
    # ============ SAMPLE NOTIFICATION TEMPLATES ============
    add_heading_styled(doc, 'Appendix B – Sample Notification Templates', 1)
    
    add_heading_styled(doc, 'Loan Request Sent to Lender', 2)
    doc.add_paragraph(
        'SMS: "Loan request LN-12345: 07XXXXXXXX requests Ksh 5,000. View details and approve in app."'
    )
    
    add_heading_styled(doc, 'Loan Disbursed', 2)
    doc.add_paragraph(
        'SMS Borrower: "Ksh 5,000 received from 07YYYYYYYY. Loan LN-12345 initialized. Balance Ksh 5,000."'
    )
    
    add_heading_styled(doc, 'Repayment Deducted', 2)
    doc.add_paragraph(
        'SMS Borrower: "Ksh 200 deducted for loan repayment. Balance Ksh 4,800."'
    )
    doc.add_paragraph(
        'SMS Lender: "Ksh 200 received from 07XXXXXXXX. Balance Ksh 4,800."'
    )
    
    add_heading_styled(doc, 'Dispute Acknowledgement', 2)
    doc.add_paragraph(
        'SMS: "Dispute ID D-98765 received. We will respond within 24 hours."'
    )
    
    doc.add_page_break()
    
    # ============ SAMPLE AGREEMENT ============
    add_heading_styled(doc, 'Appendix C – Sample Signed Agreement PDF Text', 1)
    
    add_heading_styled(doc, 'Peer-to-Peer Loan Agreement', 2)
    
    agreement_sections = [
        ('Agreement ID', 'LN-12345'),
        ('Lender', 'Name, Phone 07YYYYYYYY'),
        ('Borrower', 'Name, Phone 07XXXXXXXX'),
        ('Principal', 'Ksh 5,000'),
        ('Disbursement Date', '2026-03-01'),
        ('Repayment Method', 'Fixed Ksh 200 per trigger'),
        ('Repayment Start Date', '2026-03-05'),
    ]
    
    for section, value in agreement_sections:
        p = doc.add_paragraph()
        run = p.add_run(section)
        run.bold = True
        p.add_run(f': {value}')
    
    add_heading_styled(doc, 'Clauses', 2)
    clauses = [
        'Borrower consents that repayments will be made via M-PESA user initiated prompts when incoming funds meet the configured threshold',
        'Borrower acknowledges clause details and accepts penalties as described',
        'Penalties: grace period 7 days; late fee flat Ksh 200 per missed installment; total penalties capped at 50% of principal',
        'Dispute process described and contact details provided',
    ]
    for i, clause in enumerate(clauses, 1):
        doc.add_paragraph(f'Clause {i}: {clause}', style='List Bullet')
    
    add_heading_styled(doc, 'Consent Record', 2)
    doc.add_paragraph('Clause ticks with timestamps and device metadata')
    
    add_heading_styled(doc, 'Signature', 2)
    doc.add_paragraph('Borrower electronic signature metadata or USSD acceptance record')
    doc.add_paragraph('Lender electronic signature metadata or USSD acceptance record')
    
    add_heading_styled(doc, 'Transaction Evidence', 2)
    doc.add_paragraph('Disbursement reference M-PESA TRX: M-PESA-TRX-001')
    doc.add_paragraph('Repayment references: list of M-PESA transaction refs')
    
    add_heading_styled(doc, 'Footer', 2)
    doc.add_paragraph(
        '"This agreement is governed by Kenyan law. For questions contact support."'
    )
    
    doc.add_page_break()
    
    # ============ END ============
    final = doc.add_paragraph('END OF PROPOSAL DOCUMENT')
    final.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    final.runs[0].bold = True
    
    # Save document
    output_path = 'P2P_DEBT_MANAGEMENT_SYSTEM_PROPOSAL_v1.0.docx'
    doc.save(output_path)
    print('Word document created successfully: {}'.format(output_path))
    print('   Location: {}'.format(output_path))
    print('   Status: Ready for use')

if __name__ == '__main__':
    create_word_document()
